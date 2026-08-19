from pathlib import Path
from loguru import logger
from pypdf import PdfReader
from docx import Document as DocxDocument
import os


class TextExtractionService:

    SUPPORTED_FORMATS = {
        "text/plain": ["txt", "md"],
        "application/pdf": ["pdf"],
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ["docx"],
    }

    def extract_text(self, file_path: str, content_type: str) -> str:
        logger.info(
            "Extracting text from file",
            file_path=file_path,
            content_type=content_type,
        )
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = self._get_extension(file_path)
        logger.debug("Detected file extension", extension=extension)

        if content_type == "text/plain" or extension in ["txt", "md"]:
            result = self._extract_text_file(file_path)
        elif content_type == "application/pdf" or extension == "pdf":
            result = self._extract_pdf(file_path)
        elif "wordprocessingml" in content_type or extension == "docx":
            result = self._extract_docx(file_path)
        else:
            raise ValueError(
                f"Unsupported file format: {content_type} (extension: {extension}). "
                "Supported formats: .txt, .md, .pdf, .docx"
            )

        logger.info(
            "Text extraction complete",
            file_path=file_path,
            extracted_chars=len(result),
        )
        return result

    def _extract_text_file(self, file_path: str) -> str:
        logger.debug("Reading plain-text file", file_path=file_path)
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            logger.debug("Text file read (utf-8)", chars=len(text))
        except UnicodeDecodeError:
            logger.warning("utf-8 decode failed, falling back to latin-1", file_path=file_path)
            with open(file_path, "r", encoding="latin-1") as f:
                text = f.read()
            logger.debug("Text file read (latin-1)", chars=len(text))
        return text.strip()

    def _extract_pdf(self, file_path: str) -> str:
        logger.debug("Extracting text from PDF", file_path=file_path)
        try:
            reader = PdfReader(file_path)
            page_count = len(reader.pages)
            logger.info("PDF opened", page_count=page_count)
            full_text = "\n\n".join(
                text
                for page in reader.pages
                if (text := page.extract_text())
            )
            logger.info(
                "PDF text extracted",
                page_count=page_count,
                extracted_chars=len(full_text),
            )
            return full_text.strip()
        except Exception as e:
            logger.error("PDF extraction failed", file_path=file_path, error=str(e))
            raise Exception(f"Failed to extract text from PDF: {e}")

    def _extract_docx(self, file_path: str) -> str:
        logger.debug("Extracting text from DOCX", file_path=file_path)
        try:
            doc = DocxDocument(file_path)
            text_parts = []

            # Paragraphs
            para_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            text_parts.extend(para_texts)
            logger.debug("DOCX paragraphs extracted", count=len(para_texts))

            # Tables
            table_cells = []
            for table in doc.tables:
                for row in table.rows:
                    table_cells.extend(
                        cell.text.strip()
                        for cell in row.cells
                        if cell.text.strip()
                    )
            text_parts.extend(table_cells)
            logger.debug("DOCX table cells extracted", count=len(table_cells))

            result = "\n\n".join(text_parts)
            logger.info("DOCX extraction complete", extracted_chars=len(result))
            return result
        except Exception as e:
            logger.error("DOCX extraction failed", file_path=file_path, error=str(e))
            raise Exception(f"Failed to extract text from DOCX: {e}") from e

    def is_supported(self, content_type: str, file_path: str | Path) -> bool:
        extension = self._get_extension(file_path=file_path)
        logger.debug(
            "Checking file support",
            content_type=content_type,
            extension=extension,
        )
        if content_type in self.SUPPORTED_FORMATS:
            logger.debug("Supported via content_type", content_type=content_type)
            return True
        for mime_type, extensions in self.SUPPORTED_FORMATS.items():
            if extension in extensions:
                logger.debug("Supported via extension", extension=extension)
                return True
        logger.warning(
            "Unsupported file type",
            content_type=content_type,
            extension=extension,
        )
        return False

    def _get_extension(self, file_path: str | Path) -> str:
        return Path(file_path).suffix.lstrip(".")
